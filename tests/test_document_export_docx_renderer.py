import pytest
from docx import Document as DocxDocument

from app.schemas.document_export import (
    BulletListBlock,
    DocumentSpec,
    DocumentStyle,
    HeadingBlock,
    NumberedListBlock,
    ParagraphBlock,
    TableBlock,
)
from app.services.document_export_service import DocxRenderer


@pytest.mark.document_export
def test_docx_renderer_creates_docx_with_style_and_blocks(tmp_path):
    spec = DocumentSpec(
        title="Styled Legal Memo",
        style=DocumentStyle(
            font_family="Times New Roman",
            font_size_pt=12,
            line_spacing=1.5,
            page_size="A4",
        ),
        blocks=[
            HeadingBlock(id="h_1", level=1, text="Executive Summary"),
            ParagraphBlock(id="p_1", text="This is the legal analysis."),
            BulletListBlock(id="b_1", items=["First issue", "Second issue"]),
            NumberedListBlock(id="n_1", items=["Step one", "Step two"]),
            TableBlock(
                id="t_1",
                caption="Remedies Table",
                columns=["Remedy", "Basis"],
                rows=[
                    ["Compensation", "Employment Act"],
                    ["Reinstatement", "Employment Act"],
                ],
            ),
        ],
    )

    out_path = tmp_path / "memo.docx"

    rendered_path = DocxRenderer().render(spec, out_path)

    assert rendered_path.exists()

    doc = DocxDocument(str(rendered_path))

    assert doc.core_properties.title == "Styled Legal Memo"
    assert doc.styles["Normal"].font.name == "Times New Roman"
    assert doc.styles["Normal"].font.size.pt == 12
    assert len(doc.tables) == 1

    all_text = "\n".join(p.text for p in doc.paragraphs)

    assert "Executive Summary" in all_text
    assert "This is the legal analysis." in all_text
    assert "First issue" in all_text
    assert "Step one" in all_text
