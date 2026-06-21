"""
Tool Executor

Executes named actions from the semantic tool registry.
The runtime calls this layer when the agent selects an action.

Key design:
- Each tool is a named function
- Tools return structured observations (never empty strings or `[]`)
- Tool execution is separate from tool selection (agent decides, runtime executes)
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any, Callable

from app.config import Settings
from app.brave_search import BraveSearchClient
from app.browser_fetch import BrowserFetcher
from app.tools.kenya_law_client import KenyaLawClient
from app.tools.kenya_law_research_tool import KenyaLawResearchTool, KenyaLawRequest


# ---------------------------------------------------------------------------
# Structured observation types
# ---------------------------------------------------------------------------

@dataclass
class ToolObservation:
    """Structured result from any tool execution."""
    tool_name: str
    status: str
    message: str
    candidates: list[dict[str, Any]] = field(default_factory=list)
    error_type: str | None = None
    error_message: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "tool_name": self.tool_name,
            "status": self.status,
            "message": self.message,
            "candidate_count": len(self.candidates),
            "candidates": self.candidates[:5],  # Limit for prompt context
            "error_type": self.error_type,
            "error_message": self.error_message,
            "metadata": self.metadata,
        }


@dataclass
class ReadObservation:
    """Structured result from any fetch/read operation."""
    status: str
    url: str
    final_url: str | None = None
    content_type: str = ""
    title: str = ""
    text_excerpt: str = ""
    full_text: str = ""
    extraction_quality: str = ""
    chars_extracted: int = 0
    error_type: str | None = None
    error_message: str | None = None
    attempts: list[dict[str, Any]] = field(default_factory=list)
    suggested_alternatives: list[str] = field(default_factory=list)

    # Page structure for browser navigation — populated by html_read_browser_fetch
    links: list[dict[str, str]] = field(default_factory=list)
    forms: list[dict[str, Any]] = field(default_factory=list)
    page_narrative: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "status": self.status,
            "url": self.url,
            "final_url": self.final_url,
            "content_type": self.content_type,
            "title": self.title,
            "text_excerpt": self.text_excerpt[:500],
            "extraction_quality": self.extraction_quality,
            "chars_extracted": self.chars_extracted,
            "error_type": self.error_type,
            "error_message": self.error_message,
            "links": self.links[:20] if self.links else [],
            "forms": self.forms[:5] if self.forms else [],
            "page_narrative": self.page_narrative[:300] if self.page_narrative else "",
        }


# ---------------------------------------------------------------------------
# Tool Executor
# ---------------------------------------------------------------------------

class ToolExecutor:
    """
    Executes named tools and returns structured observations.

    The agent selects a tool name + parameters. The executor:
    1. Validates the tool exists
    2. Validates inputs
    3. Executes the tool
    4. Normalizes the observation
    5. Returns a structured ToolObservation or ReadObservation
    """

    def __init__(self, settings: Settings, llm=None) -> None:
        self.settings = settings
        self._llm = llm  # passed through to KenyaLawResearchTool

        # Initialize clients
        try:
            self.brave = BraveSearchClient(settings)
        except Exception as e:
            self.brave = None
            self._brave_error = str(e)

        try:
            self.browser = BrowserFetcher(settings)
        except Exception as e:
            self.browser = None
            self._browser_error = str(e)

        try:
            self.kenya_law = KenyaLawClient()
        except Exception:
            self.kenya_law = KenyaLawClient()

        # Unified Kenya Law Research Tool — LLM-guided pipeline
        self.kenya_law_research_tool = KenyaLawResearchTool(
            llm=self._llm,
            brave=self.brave,
        )

        self._brave_error: str = ""

    def list_tools(self) -> list[str]:
        """Return list of available tool names."""
        return [
            "kenya_law_research",        # Unified Kenya Law tool (preferred)
            "kenya_law_judgment_search",  # Legacy — kept for backward compat
            "kenya_law_legislation_search",
            "kenya_law_case_resolve",
            "case_specific_search",
            "official_kenya_domain_search",
            "brave_search",
            "general_web_fetch",
            "kenya_law_read",
            "akn_xml_read",
            "pdf_read",
            "docx_read",
            "html_read_browser_fetch",
        ]

    async def execute(self, tool_name: str, params: dict[str, Any]) -> ToolObservation | ReadObservation:
        """Execute a named tool with given parameters."""
        executor_map: dict[str, Callable] = {
            "kenya_law_research": self._kenya_law_research,  # Unified tool
            "kenya_law_judgment_search": self._kenya_law_judgment_search,
            "kenya_law_legislation_search": self._kenya_law_legislation_search,
            "kenya_law_case_resolve": self._kenya_law_case_resolve,
            "case_specific_search": self._case_specific_search,
            "official_kenya_domain_search": self._official_kenya_domain_search,
            "brave_search": self._brave_search,
            "general_web_fetch": self._general_web_fetch,
            "kenya_law_read": self._kenya_law_read,
            "akn_xml_read": self._akn_xml_read,
            "pdf_read": self._pdf_read,
            "docx_read": self._docx_read,
            "html_read_browser_fetch": self._html_read_browser_fetch,
        }

        executor = executor_map.get(tool_name)
        if executor is None:
            return ToolObservation(
                tool_name=tool_name,
                status="unknown_tool",
                message=f"Tool '{tool_name}' is not registered in the executor.",
                error_type="unknown_tool",
            )

        return await executor(params)

    # -----------------------------------------------------------------------
    # Kenya Law Judgment Search
    # -----------------------------------------------------------------------

    async def _kenya_law_judgment_search(self, params: dict[str, Any]) -> ToolObservation:
        query = params.get("query", "")
        if not query:
            return ToolObservation(
                tool_name="kenya_law_judgment_search",
                status="invalid_input",
                message="No search query provided.",
                error_type="missing_query",
            )

        try:
            candidates = await self.kenya_law.search_judgments(query)
        except Exception as e:
            return ToolObservation(
                tool_name="kenya_law_judgment_search",
                status="network_error",
                message=f"Kenya Law judgment search failed: {e}",
                error_type="network_error",
                error_message=str(e),
            )

        if not candidates:
            return ToolObservation(
                tool_name="kenya_law_judgment_search",
                status="no_results",
                message="Kenya Law judgment search returned no results for the query.",
                metadata={"query": query},
            )

        return ToolObservation(
            tool_name="kenya_law_judgment_search",
            status="success_with_candidates",
            message=f"Found {len(candidates)} judgment candidates.",
            candidates=[
                {
                    "source_id": c.source_id,
                    "title": c.title,
                    "url": c.url,
                    "snippet": c.snippet[:200],
                    "score": c.confidence,
                }
                for c in candidates[:10]
            ],
            metadata={"query": query, "candidate_count": len(candidates)},
        )

    # -----------------------------------------------------------------------
    # Kenya Law Legislation Search
    # -----------------------------------------------------------------------

    async def _kenya_law_legislation_search(self, params: dict[str, Any]) -> ToolObservation:
        query = params.get("query", "")
        if not query:
            return ToolObservation(
                tool_name="kenya_law_legislation_search",
                status="invalid_input",
                message="No search query provided.",
                error_type="missing_query",
            )

        try:
            candidates = await self.kenya_law.search_legislation(query)
        except Exception as e:
            return ToolObservation(
                tool_name="kenya_law_legislation_search",
                status="network_error",
                message=f"Kenya Law legislation search failed: {e}",
                error_type="network_error",
                error_message=str(e),
            )

        if not candidates:
            return ToolObservation(
                tool_name="kenya_law_legislation_search",
                status="no_results",
                message="Kenya Law legislation search returned no results.",
                metadata={"query": query},
            )

        return ToolObservation(
            tool_name="kenya_law_legislation_search",
            status="success_with_candidates",
            message=f"Found {len(candidates)} legislation candidates.",
            candidates=[
                {
                    "source_id": c.source_id,
                    "title": c.title,
                    "url": c.url,
                    "snippet": c.snippet[:200],
                }
                for c in candidates[:10]
            ],
            metadata={"query": query, "candidate_count": len(candidates)},
        )

    # -----------------------------------------------------------------------
    # Kenya Law Case Resolve
    # -----------------------------------------------------------------------

    async def _kenya_law_case_resolve(self, params: dict[str, Any]) -> ToolObservation:
        citation = params.get("citation", "")
        court = params.get("court", "")
        year = params.get("year", "")
        number = params.get("number", "")

        if not citation and not (court and year and number):
            return ToolObservation(
                tool_name="kenya_law_case_resolve",
                status="invalid_input",
                message="Either a citation string or court+year+number is required.",
                error_type="missing_parameters",
            )

        # Attempt to resolve using the client
        # Build a search from the citation metadata
        search_terms = citation or f"{court} {year} {number}"

        try:
            # Search judgments for the specific citation
            candidates = await self.kenya_law.search_judgments(search_terms)
        except Exception as e:
            return ToolObservation(
                tool_name="kenya_law_case_resolve",
                status="network_error",
                message=f"Kenya Law case resolution failed: {e}",
                error_type="network_error",
                error_message=str(e),
            )

        if not candidates:
            # Try to reconstruct an AKN URL
            reconstructed_url = None
            if court and year and number:
                court_code = court.upper()
                reconstructed_url = f"https://new.kenyalaw.org/akn/ke/judgment/ke{ court_code.lower() }/{year}/{number}"
                return ToolObservation(
                    tool_name="kenya_law_case_resolve",
                    status="generated_unverified",
                    message=f"Could not find exact match. Generated unverified URL: {reconstructed_url}",
                    metadata={
                        "generated_url": reconstructed_url,
                        "citation": citation,
                        "court": court,
                        "year": year,
                        "number": number,
                    },
                )

            return ToolObservation(
                tool_name="kenya_law_case_resolve",
                status="not_found",
                message="Could not resolve the case from the provided citation metadata.",
                metadata={"citation": citation},
            )

        return ToolObservation(
            tool_name="kenya_law_case_resolve",
            status="resolved_verified",
            message=f"Resolved {len(candidates)} candidate(s) for the citation.",
            candidates=[
                {
                    "source_id": c.source_id,
                    "title": c.title,
                    "url": c.url,
                }
                for c in candidates[:5]
            ],
            metadata={"citation": citation, "candidate_count": len(candidates)},
        )

    # -----------------------------------------------------------------------
    # Case-Specific Search
    # -----------------------------------------------------------------------

    async def _case_specific_search(self, params: dict[str, Any]) -> ToolObservation:
        query = params.get("query", "")
        parties = params.get("parties", "")
        case_number = params.get("case_number", "")
        issues = params.get("issues", [])

        if not query and not parties and not case_number:
            return ToolObservation(
                tool_name="case_specific_search",
                status="invalid_input",
                message="Either a query, parties, or case number is required.",
                error_type="missing_parameters",
            )

        search_text = query or f"{parties} {case_number} {' '.join(issues)}".strip()

        try:
            candidates = await self.kenya_law.search_judgments(search_text)
        except Exception as e:
            return ToolObservation(
                tool_name="case_specific_search",
                status="network_error",
                message=f"Case-specific search failed: {e}",
                error_type="network_error",
                error_message=str(e),
            )

        if not candidates:
            return ToolObservation(
                tool_name="case_specific_search",
                status="no_results",
                message="Case-specific search returned no results.",
                metadata={"query": search_text},
            )

        return ToolObservation(
            tool_name="case_specific_search",
            status="success_with_candidates",
            message=f"Found {len(candidates)} candidate(s).",
            candidates=[
                {
                    "source_id": c.source_id,
                    "title": c.title,
                    "url": c.url,
                    "score": c.confidence,
                }
                for c in candidates[:10]
            ],
            metadata={"query": search_text, "candidate_count": len(candidates)},
        )

    # -----------------------------------------------------------------------
    # Kenya Law Research — unified tool (replaces fragmented KL tools)
    # -----------------------------------------------------------------------

    async def _kenya_law_research(self, params: dict[str, Any]) -> ToolObservation:
        """Unified Kenya Law research capability.

        Routes to the new KenyaLawResearchTool which handles:
        - interpret → discover → rank → fetch → parse → verify
        - legislation, case law, gazette, AKN URLs, PDF/DOCX
        """
        query = params.get("query", "")
        if not query and not params.get("known_url"):
            return ToolObservation(
                tool_name="kenya_law_research",
                status="invalid_input",
                message="Either a query or a known_url is required.",
                error_type="missing_input",
            )

        request = KenyaLawRequest(
            query=query,
            kind=params.get("kind", "auto"),
            operation=params.get("operation", "search_and_fetch"),
            known_url=params.get("known_url"),
            citation=params.get("citation"),
            case_number=params.get("case_number"),
            court=params.get("court"),
            section=params.get("section"),
            article=params.get("article"),
            date_from=params.get("date_from"),
            date_to=params.get("date_to"),
            max_candidates=params.get("max_candidates", 10),
            read_top_n=params.get("read_top_n", 3),
        )

        try:
            result = await self.kenya_law_research_tool.research(request)
        except Exception as e:
            return ToolObservation(
                tool_name="kenya_law_research",
                status="error",
                message=f"Kenya Law research failed: {e}",
                error_type="research_error",
                error_message=str(e),
            )

        # Convert KenyaLawResult to ToolObservation
        d = result.to_tool_observation_dict()

        # Build candidates list for evidence ledger
        candidates_list = []
        for c in result.candidates[:10]:
            candidates_list.append({
                "title": c.title,
                "url": c.url,
                "snippet": (c.snippet or "")[:200],
                "score": c.score,
                "source_id": f"KL_{hash(c.url) % 10000:04d}",
            })

        return ToolObservation(
            tool_name="kenya_law_research",
            status=d["status"],
            message=d["message"],
            candidates=candidates_list if candidates_list else [],
            metadata={
                "intent_kind": d.get("intent_kind"),
                "document_title": d.get("selected_document_title"),
                "document_url": d.get("selected_document_url"),
                "document_text": result.selected_document_text[:50000] if result.selected_document_text else None,
                "provision_text": result.selected_document_provision,
                "provision_extracted": bool(result.selected_document_provision),
                "chars_extracted": len(result.selected_document_text) if result.selected_document_text else 0,
                "retrieval_attempts": d.get("retrieval_attempts", []),
                "suggested_next": d.get("suggested_next_actions", []),
            },
        )

    # -----------------------------------------------------------------------
    # Official Kenya Domain Search
    # -----------------------------------------------------------------------

    async def _official_kenya_domain_search(self, params: dict[str, Any]) -> ToolObservation:
        query = params.get("query", "")
        if not query:
            return ToolObservation(
                tool_name="official_kenya_domain_search",
                status="invalid_input",
                message="No search query provided.",
                error_type="missing_query",
            )

        if not self.brave:
            return ToolObservation(
                tool_name="official_kenya_domain_search",
                status="api_error",
                message=f"Brave Search API is not available: {self._brave_error}",
                error_type="brave_not_initialized",
            )

        # Add site filters for official Kenyan domains
        domains = [
            "site:judiciary.go.ke",
            "site:parliament.go.ke",
            "site:klrc.go.ke",
            "site:kenyalaw.org",
        ]
        domain_query = f"{query} ({' OR '.join(domains)})"

        try:
            results = await self.brave.search(domain_query)
        except Exception as e:
            return ToolObservation(
                tool_name="official_kenya_domain_search",
                status="network_error",
                message=f"Official Kenya domain search failed: {e}",
                error_type="network_error",
                error_message=str(e),
            )

        if not results:
            return ToolObservation(
                tool_name="official_kenya_domain_search",
                status="no_results",
                message="No results from official Kenyan domains.",
                metadata={"query": query},
            )

        return ToolObservation(
            tool_name="official_kenya_domain_search",
            status="success_with_candidates",
            message=f"Found {len(results)} results from official Kenyan domains.",
            candidates=[
                {
                    "source_id": getattr(r, "id", f"src_{i}"),
                    "title": getattr(r, "title", "Untitled"),
                    "url": getattr(r, "url", ""),
                    "snippet": getattr(r, "snippet", "")[:200],
                }
                for i, r in enumerate(results[:10])
            ],
            metadata={"query": query, "candidate_count": len(results)},
        )

    # -----------------------------------------------------------------------
    # Brave Search
    # -----------------------------------------------------------------------

    async def _brave_search(self, params: dict[str, Any]) -> ToolObservation:
        query = params.get("query", "")
        if not query:
            return ToolObservation(
                tool_name="brave_search",
                status="invalid_input",
                message="No search query provided.",
                error_type="missing_query",
            )

        if not self.brave:
            return ToolObservation(
                tool_name="brave_search",
                status="api_error",
                message=f"Brave Search API is not available: {self._brave_error}",
                error_type="brave_not_initialized",
            )

        try:
            results = await self.brave.search(query)
        except Exception as e:
            return ToolObservation(
                tool_name="brave_search",
                status="network_error",
                message=f"Brave search failed: {e}",
                error_type="network_error",
                error_message=str(e),
            )

        if not results:
            return ToolObservation(
                tool_name="brave_search",
                status="no_results",
                message="Brave search returned no results.",
                metadata={"query": query},
            )

        return ToolObservation(
            tool_name="brave_search",
            status="success_with_candidates",
            message=f"Found {len(results)} results via Brave.",
            candidates=[
                {
                    "source_id": getattr(r, "id", f"src_{i}"),
                    "title": getattr(r, "title", "Untitled"),
                    "url": getattr(r, "url", ""),
                    "snippet": getattr(r, "snippet", "")[:200],
                }
                for i, r in enumerate(results[:10])
            ],
            metadata={"query": query, "candidate_count": len(results)},
        )

    # -----------------------------------------------------------------------
    # General Web Fetch
    # -----------------------------------------------------------------------

    async def _general_web_fetch(self, params: dict[str, Any]) -> ReadObservation:
        url = params.get("url", "")
        if not url:
            return ReadObservation(
                status="invalid_input",
                url="",
                error_type="missing_url",
                error_message="No URL provided for fetching.",
            )

        attempts: list[dict[str, Any]] = []

        # Attempt 1: Direct HTTP fetch
        try:
            import httpx
            _headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "en-GB,en;q=0.9",
            }
            async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
                resp = await client.get(url, headers=_headers)
                attempts.append({
                    "method": "http_fetch",
                    "status_code": resp.status_code,
                    "content_type": resp.headers.get("content-type", ""),
                })

                if resp.status_code == 200:
                    content_type = resp.headers.get("content-type", "").lower()
                    url_lower = url.lower()

                    # PDF detection: delegate to pypdf instead of treating bytes as HTML
                    if "pdf" in content_type or url_lower.endswith(".pdf"):
                        from pypdf import PdfReader
                        try:
                            reader = PdfReader(io.BytesIO(resp.content))
                            pdf_parts = []
                            for page in reader.pages[:20]:
                                t = page.extract_text()
                                if t:
                                    pdf_parts.append(t)
                            pdf_text = "\n".join(pdf_parts).strip()
                            if pdf_text:
                                return ReadObservation(
                                    status="read_success",
                                    url=url,
                                    final_url=str(resp.url),
                                    content_type=content_type,
                                    title="",
                                    text_excerpt=pdf_text[:500],
                                    full_text=pdf_text,
                                    extraction_quality="good" if len(pdf_text) > 500 else "minimal",
                                    chars_extracted=len(pdf_text),
                                    attempts=attempts,
                                )
                        except Exception:
                            attempts.append({"method": "pdf_parse", "error": "pypdf_extraction_failed"})
                        # PDF was unreadable — fall through to remaining attempts

                    raw_html = resp.text
                    # Try to extract meaningful text
                    title_match = re.search(r'<title[^>]*>(.*?)</title>', raw_html, re.IGNORECASE | re.DOTALL)
                    title = title_match.group(1).strip() if title_match else ""

                    # Try document-content div first
                    doc_match = re.search(
                        r'<div[^>]*id=[\'"]document-content[\'"][^>]*>(.*?)</div>',
                        raw_html, re.DOTALL | re.IGNORECASE
                    )
                    if doc_match:
                        text = re.sub(r'<[^>]+>', ' ', doc_match.group(1))
                    else:
                        text = re.sub(r'<[^>]+>', ' ', raw_html)
                    text = re.sub(r'\s+', ' ', text).strip()

                    if len(text) > 50:
                        return ReadObservation(
                            status="read_success",
                            url=url,
                            final_url=str(resp.url),
                            content_type=resp.headers.get("content-type", ""),
                            title=title,
                            text_excerpt=text[:500],
                            full_text=text,
                            extraction_quality="good" if len(text) > 500 else "minimal",
                            chars_extracted=len(text),
                            attempts=attempts,
                        )

        except httpx.TimeoutException:
            attempts.append({"method": "http_fetch", "error": "timeout"})
        except httpx.HTTPStatusError as e:
            attempts.append({"method": "http_fetch", "error": f"http_{e.response.status_code}"})
        except Exception as e:
            attempts.append({"method": "http_fetch", "error": str(e)[:100]})

        # Attempt 2: Try browser fetch
        if self.browser:
            try:
                text = await self.browser.fetch_text(url)
                attempts.append({"method": "browser_fetch", "status": "attempted", "chars": len(text) if text else 0})
                if text and len(text) > 50:
                    return ReadObservation(
                        status="read_success",
                        url=url,
                        final_url=url,
                        content_type="text/html",
                        title="",
                        text_excerpt=text[:500],
                        full_text=text,
                        extraction_quality="good" if len(text) > 500 else "minimal",
                        chars_extracted=len(text),
                        attempts=attempts,
                    )
            except Exception as e:
                attempts.append({"method": "browser_fetch", "error": str(e)[:100]})

        # All attempts failed
        error_msg = "All fetch methods failed to extract readable content."
        if attempts:
            last_attempt = attempts[-1]
            error_type = last_attempt.get("error", "extraction_failed")
            error_msg = f"Fetch failed. Last attempt: {last_attempt.get('method', 'unknown')} - {last_attempt.get('error', 'unknown error')}"

        return ReadObservation(
            status="unreadable",
            url=url,
            error_type=error_type,
            error_message=error_msg,
            attempts=attempts,
        )

    # -----------------------------------------------------------------------
    # Kenya Law Read
    # -----------------------------------------------------------------------

    async def _kenya_law_read(self, params: dict[str, Any]) -> ReadObservation:
        url = params.get("url", "")
        if not url:
            return ReadObservation(
                status="invalid_input",
                url="",
                error_type="missing_url",
                error_message="No URL provided.",
            )

        import httpx
        from bs4 import BeautifulSoup

        _headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        }

        attempts: list[dict[str, Any]] = []
        url_lower = url.lower()

        try:
            async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:

                # Ladder 1: HTML body extraction
                resp = await client.get(url, headers=_headers)
                attempts.append({
                    "method": "html_body",
                    "status_code": resp.status_code,
                })

                if resp.status_code == 200:
                    soup = BeautifulSoup(resp.text, "html.parser")

                    # Look for document-content div
                    content_div = soup.find("div", id=re.compile(r"document-content", re.I))
                    if content_div:
                        text = content_div.get_text(" ", strip=True)
                        title_tag = soup.find("title")
                        title = title_tag.get_text(strip=True) if title_tag else ""

                        if len(text) > 100:
                            return ReadObservation(
                                status="read_success",
                                url=url,
                                final_url=str(resp.url),
                                title=title,
                                text_excerpt=text[:500],
                                full_text=text,
                                extraction_quality="good" if len(text) > 500 else "minimal",
                                chars_extracted=len(text),
                                attempts=attempts,
                            )
                        attempts.append({
                            "method": "html_body",
                            "status": "failed",
                            "reason": "body_container_not_found_or_empty",
                        })

                    else:
                        # Try full page text
                        text = soup.get_text(" ", strip=True)
                        if len(text) > 200:
                            return ReadObservation(
                                status="read_success",
                                url=url,
                                title=soup.find("title").get_text(strip=True) if soup.find("title") else "",
                                text_excerpt=text[:500],
                                full_text=text,
                                extraction_quality="minimal",
                                chars_extracted=len(text),
                                attempts=attempts,
                            )
                        attempts.append({
                            "method": "html_body",
                            "status": "failed",
                            "reason": "insufficient_text_extracted",
                        })

                # Ladder 2: AKN XML
                if "/akn/" in url_lower:
                    akn_url = url.rstrip("/")
                    try:
                        akn_resp = await client.get(
                            akn_url,
                            headers={**_headers, "Accept": "application/xml,text/xml,*/*"},
                        )
                        attempts.append({
                            "method": "akn_xml",
                            "status_code": akn_resp.status_code,
                        })
                        if akn_resp.status_code == 200 and len(akn_resp.text) > 200:
                            # Extract text from XML
                            xml_soup = BeautifulSoup(akn_resp.text, "xml")
                            xml_text = xml_soup.get_text(" ", strip=True) if xml_soup else akn_resp.text
                            if len(xml_text) > 200:
                                return ReadObservation(
                                    status="read_success",
                                    url=url,
                                    final_url=str(akn_resp.url),
                                    title="",
                                    text_excerpt=xml_text[:500],
                                    full_text=xml_text,
                                    extraction_quality="good",
                                    chars_extracted=len(xml_text),
                                    attempts=attempts,
                                )
                    except Exception as e:
                        attempts.append({"method": "akn_xml", "error": str(e)[:100]})

                # Ladder 3: PDF download
                if "/akn/" in url_lower:
                    pdf_url = f"{url.rstrip('/')}/source.pdf"
                    try:
                        pdf_resp = await client.get(pdf_url, headers=_headers)
                        attempts.append({
                            "method": "pdf_download",
                            "status_code": pdf_resp.status_code,
                        })
                        if pdf_resp.status_code == 200 and len(pdf_resp.content) > 500:
                            try:
                                from pypdf import PdfReader
                                reader = PdfReader(io.BytesIO(pdf_resp.content))
                                pdf_text = []
                                for page in reader.pages:
                                    t = page.extract_text()
                                    if t:
                                        pdf_text.append(t)
                                text = "\n".join(pdf_text)
                                if text.strip():
                                    return ReadObservation(
                                        status="read_success",
                                        url=url,
                                        title="",
                                        text_excerpt=text[:500],
                                        full_text=text,
                                        extraction_quality="good",
                                        chars_extracted=len(text),
                                        attempts=attempts,
                                    )
                            except Exception as e:
                                attempts.append({"method": "pdf_parse", "error": str(e)[:100]})
                    except Exception as e:
                        attempts.append({"method": "pdf_download", "error": str(e)[:100]})

                # Ladder 4: DOCX download
                if "/akn/" in url_lower:
                    docx_url = f"{url.rstrip('/')}/source"
                    try:
                        docx_resp = await client.get(docx_url, headers=_headers)
                        attempts.append({
                            "method": "docx_download",
                            "status_code": docx_resp.status_code,
                        })
                        if docx_resp.status_code == 200 and len(docx_resp.content) > 200:
                            try:
                                from docx import Document
                                doc = Document(io.BytesIO(docx_resp.content))
                                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                                if text.strip():
                                    return ReadObservation(
                                        status="read_success",
                                        url=url,
                                        title="",
                                        text_excerpt=text[:500],
                                        full_text=text,
                                        extraction_quality="good",
                                        chars_extracted=len(text),
                                        attempts=attempts,
                                    )
                            except Exception as e:
                                attempts.append({"method": "docx_parse", "error": str(e)[:100]})
                    except Exception as e:
                        attempts.append({"method": "docx_download", "error": str(e)[:100]})

        except httpx.TimeoutException:
            attempts.append({"method": "http_fetch", "error": "timeout"})
        except Exception as e:
            attempts.append({"method": "http_fetch", "error": str(e)[:100]})

        return ReadObservation(
            status="unreadable",
            url=url,
            error_type="all_methods_failed",
            error_message="Could not extract readable text from Kenya Law page via any available method.",
            attempts=attempts,
        )

    # -----------------------------------------------------------------------
    # AKN XML Read
    # -----------------------------------------------------------------------

    async def _akn_xml_read(self, params: dict[str, Any]) -> ReadObservation:
        url = params.get("url", "")
        if not url:
            return ReadObservation(
                status="invalid_input",
                url="",
                error_type="missing_url",
                error_message="No URL provided.",
            )

        try:
            xml_text = await self.kenya_law.fetch_akn_xml(url)
        except Exception as e:
            return ReadObservation(
                status="parser_error",
                url=url,
                error_type="fetch_failed",
                error_message=f"AKN XML fetch failed: {e}",
            )

        if not xml_text:
            return ReadObservation(
                status="not_found",
                url=url,
                error_type="empty_response",
                error_message="AKN XML fetch returned empty response.",
            )

        # Try to extract text from XML
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(xml_text, "xml")
            extracted = soup.get_text(" ", strip=True)
        except Exception:
            extracted = xml_text[:500]

        return ReadObservation(
            status="read_success",
            url=url,
            content_type="application/xml",
            text_excerpt=extracted[:500],
            full_text=extracted,
            extraction_quality="good" if len(extracted) > 200 else "minimal",
            chars_extracted=len(extracted),
        )

    # -----------------------------------------------------------------------
    # PDF Read
    # -----------------------------------------------------------------------

    async def _pdf_read(self, params: dict[str, Any]) -> ReadObservation:
        url = params.get("url", "")
        if not url:
            return ReadObservation(
                status="invalid_input",
                url="",
                error_type="missing_url",
                error_message="No URL provided.",
            )

        import httpx

        _headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        }

        try:
            async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
                resp = await client.get(url, headers=_headers)

                if resp.status_code == 404:
                    return ReadObservation(
                        status="not_found",
                        url=url,
                        error_type="http_404",
                        error_message="PDF URL returned 404 Not Found.",
                    )

                if resp.status_code != 200:
                    return ReadObservation(
                        status="blocked",
                        url=url,
                        error_type=f"http_{resp.status_code}",
                        error_message=f"HTTP {resp.status_code} when fetching PDF.",
                    )

                if len(resp.content) < 200:
                    return ReadObservation(
                        status="empty_body",
                        url=url,
                        error_type="empty_pdf",
                        error_message="PDF response body was empty or too small.",
                    )

                try:
                    from pypdf import PdfReader
                    reader = PdfReader(io.BytesIO(resp.content))
                    text_parts = []
                    for page in reader.pages[:20]:  # Limit to 20 pages
                        t = page.extract_text()
                        if t:
                            text_parts.append(t)
                    text = "\n".join(text_parts)

                    if not text.strip():
                        return ReadObservation(
                            status="ocr_required",
                            url=url,
                            error_type="image_only_pdf",
                            error_message="PDF appears to contain only scanned images. OCR may be required.",
                            chars_extracted=0,
                        )

                    return ReadObservation(
                        status="read_success",
                        url=url,
                        content_type="application/pdf",
                        text_excerpt=text[:500],
                        full_text=text,
                        extraction_quality="good" if len(text) > 500 else "minimal",
                        chars_extracted=len(text),
                    )

                except ImportError:
                    return ReadObservation(
                        status="unsupported_format",
                        url=url,
                        error_type="missing_dependency",
                        error_message="PDF parsing library (pypdf) is not installed.",
                    )
                except Exception as e:
                    return ReadObservation(
                        status="parser_error",
                        url=url,
                        error_type="pdf_parse_error",
                        error_message=f"Failed to parse PDF: {e}",
                    )

        except httpx.TimeoutException:
            return ReadObservation(
                status="timeout",
                url=url,
                error_type="timeout",
                error_message="PDF fetch timed out.",
            )
        except Exception as e:
            return ReadObservation(
                status="blocked",
                url=url,
                error_type="fetch_error",
                error_message=f"PDF fetch failed: {e}",
            )

    # -----------------------------------------------------------------------
    # DOCX Read
    # -----------------------------------------------------------------------

    async def _docx_read(self, params: dict[str, Any]) -> ReadObservation:
        url = params.get("url", "")
        if not url:
            return ReadObservation(
                status="invalid_input",
                url="",
                error_type="missing_url",
                error_message="No URL provided.",
            )

        import httpx

        _headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        }

        try:
            async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
                resp = await client.get(url, headers=_headers)

                if resp.status_code != 200:
                    return ReadObservation(
                        status="not_found",
                        url=url,
                        error_type=f"http_{resp.status_code}",
                        error_message=f"HTTP {resp.status_code} when fetching DOCX.",
                    )

                try:
                    from docx import Document
                    doc = Document(io.BytesIO(resp.content))
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

                    if not text.strip():
                        return ReadObservation(
                            status="empty_body",
                            url=url,
                            error_type="empty_document",
                            error_message="DOCX file contained no readable text.",
                        )

                    return ReadObservation(
                        status="read_success",
                        url=url,
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        text_excerpt=text[:500],
                        full_text=text,
                        extraction_quality="good" if len(text) > 200 else "minimal",
                        chars_extracted=len(text),
                    )

                except ImportError:
                    return ReadObservation(
                        status="unsupported_format",
                        url=url,
                        error_type="missing_dependency",
                        error_message="DOCX parsing library (python-docx) is not installed.",
                    )
                except Exception as e:
                    return ReadObservation(
                        status="parser_error",
                        url=url,
                        error_type="docx_parse_error",
                        error_message=f"Failed to parse DOCX: {e}",
                    )

        except Exception as e:
            return ReadObservation(
                status="blocked",
                url=url,
                error_type="fetch_error",
                error_message=f"DOCX fetch failed: {e}",
            )

    # -----------------------------------------------------------------------
    # HTML Read / Browser Fetch
    # -----------------------------------------------------------------------

    async def _html_read_browser_fetch(self, params: dict[str, Any]) -> ReadObservation:
        """Fetch a URL using Playwright headless browser. Returns page text + links + forms for navigation."""
        url = params.get("url", "")
        if not url:
            return ReadObservation(
                status="invalid_input",
                url="",
                error_type="missing_url",
                error_message="No URL provided.",
            )

        attempts: list[dict[str, Any]] = []

        # Structured browser fetch — returns text + links + forms
        if self.browser:
            try:
                page = await self.browser.fetch_structured(url)
                attempts.append({
                    "method": "browser_fetch_structured",
                    "chars_extracted": len(page.visible_text) if page.visible_text else 0,
                    "links_found": len(page.links),
                    "forms_found": len(page.forms),
                })

                # Build a page narrative: what links are available for navigation
                nav_lines = []
                if page.links:
                    nav_lines.append("Links on this page:")
                    for i, link in enumerate(page.links[:25], 1):
                        nav_lines.append(f"  [{i}] {link['text'][:80]}")
                        nav_lines.append(f"      URL: {link['href'][:120]}")
                if page.forms:
                    nav_lines.append("Forms on this page:")
                    for f in page.forms[:5]:
                        fields = "; ".join(f["fields"][:5])
                        nav_lines.append(f"  Form action: {f['action']}")
                        nav_lines.append(f"  Fields: {fields}")
                page_narrative = "\n".join(nav_lines)

                if page.visible_text and len(page.visible_text) > 50:
                    return ReadObservation(
                        status="read_success",
                        url=page.url,
                        final_url=page.url,
                        title=page.title or "",
                        text_excerpt=page.visible_text[:500],
                        full_text=page.visible_text,
                        extraction_quality="good" if len(page.visible_text) > 500 else "minimal",
                        chars_extracted=len(page.visible_text),
                        attempts=attempts,
                        links=page.links[:30],
                        forms=page.forms[:5],
                        page_narrative=page_narrative,
                    )
            except Exception as e:
                attempts.append({"method": "browser_fetch_structured", "error": str(e)[:100]})

        # Fall back to legacy browser fetch (text only)
        if self.browser:
            try:
                text = await self.browser.fetch_text(url)
                attempts.append({
                    "method": "browser_fetch_legacy",
                    "chars_extracted": len(text) if text else 0,
                })
                if text and len(text) > 100:
                    return ReadObservation(
                        status="read_success",
                        url=url,
                        final_url=url,
                        content_type="text/html",
                        text_excerpt=text[:500],
                        full_text=text,
                        extraction_quality="good" if len(text) > 500 else "minimal",
                        chars_extracted=len(text),
                        attempts=attempts,
                    )
            except Exception as e:
                attempts.append({"method": "browser_fetch_legacy", "error": str(e)[:100]})

        # Fall back to general web fetch
        return await self._general_web_fetch(params)
