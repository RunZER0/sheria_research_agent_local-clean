"""
Kenya Law Research Tool — pure LLM-guided pipeline.

Code is an enabler. The LLM makes every decision:
- what to search for (including site: restrictions)
- which candidate is relevant
- what metadata to extract
- what section text/passages to return

Code only: fetches bytes, parses PDF/DOCX into text, validates JSON,
records observability.
"""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from typing import Any
from urllib.parse import urlparse, urlunparse

import httpx
from bs4 import BeautifulSoup


# ===========================================================================
# Data contracts
# ===========================================================================

@dataclass
class KenyaLawRequest:
    query: str
    kind: str = "auto"
    operation: str = "search_and_fetch"
    known_url: str | None = None
    citation: str | None = None
    case_number: str | None = None
    court: str | None = None
    section: str | int | None = None
    article: str | int | None = None
    date_from: str | None = None
    date_to: str | None = None
    max_candidates: int = 10
    read_top_n: int = 3


@dataclass
class KenyaLawCandidate:
    title: str
    url: str
    snippet: str | None
    kind: str
    discovery_method: str
    score: float = 0.0
    reasons: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


@dataclass
class RetrievalAttempt:
    method: str
    url: str
    status: str
    http_status: int | None = None
    content_type: str | None = None
    chars_extracted: int = 0
    error: str | None = None


@dataclass
class Passage:
    text: str
    matched_terms: list[str]
    score: float
    start_char: int
    end_char: int


@dataclass
class KenyaLawDocument:
    kind: str
    title: str
    original_url: str | None
    normalized_akn_url: str | None
    content_url: str | None
    metadata: dict
    full_text: str
    requested_provision_text: str | None
    relevant_passages: list[Passage]
    extraction_quality: str
    authority_level: str = "primary"


@dataclass
class KenyaLawResult:
    status: str
    message: str
    query: str
    candidates: list[KenyaLawCandidate]
    rejected_candidates: list[dict]
    retrieval_attempts: list[RetrievalAttempt]
    selected_document_title: str | None = None
    selected_document_url: str | None = None
    selected_document_provision: str | None = None
    selected_document_text: str | None = None
    suggested_next_actions: list[str] = field(default_factory=list)
    llm_audit: list[dict] = field(default_factory=list)

    def to_tool_observation_dict(self) -> dict[str, Any]:
        return {
            "status": self.status,
            "message": self.message,
            "query": self.query,
            "selected_document_title": self.selected_document_title,
            "selected_document_url": self.selected_document_url,
            "provision_extracted": bool(self.selected_document_provision),
            "candidate_count": len(self.candidates),
            "llm_calls": len(self.llm_audit),
        }


# ===========================================================================
# Pure-code helpers — no decisions, just enabling
# ===========================================================================

def normalize_akn_url(url: str) -> str:
    parsed = urlparse(url)
    path = parsed.path.rstrip("/")
    parts = path.split("/")
    if parts and "@" in parts[-1]:
        parts[-1] = parts[-1].split("@", 1)[0]
    return urlunparse(parsed._replace(path="/".join(parts), query="", fragment="")).rstrip("/")


def source_url(normalized: str) -> str:
    return normalized.rstrip("/") + "/source"


def source_pdf_url(normalized: str) -> str:
    return normalized.rstrip("/") + "/source.pdf"


_KL_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/125.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-GB,en;q=0.9",
}


async def _http_get(url: str, accept: str | None = None, as_bytes: bool = False) -> tuple[int, Any, str]:
    headers = dict(_KL_HEADERS)
    if accept:
        headers["Accept"] = accept
    async with httpx.AsyncClient(timeout=httpx.Timeout(30), follow_redirects=True) as c:
        r = await c.get(url, headers=headers)
        ct = r.headers.get("content-type", "")
        if as_bytes:
            return r.status_code, r.content, ct
        return r.status_code, r.text, ct


def _parse_bytes(body: bytes, content_type: str) -> str | None:
    """Pure code: try PDF, then DOCX, then raw text. Returns text or None."""
    ct = content_type.lower() if content_type else ""
    try:
        if "pdf" in ct or (b"%PDF" in body[:10]):
            from pypdf import PdfReader
            r = PdfReader(io.BytesIO(body))
            text = "\n".join(p.extract_text() or "" for p in r.pages[:20]).strip()
            if text:
                return text
        try:
            from docx import Document
            doc = Document(io.BytesIO(body))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
            if text:
                return text
        except Exception:
            pass
        decoded = body.decode("utf-8", errors="replace").strip()
        # Reject binary data disguised as text: if > 30% of characters are
        # replacement chars (�), the PDF was not parsed correctly and we are
        # looking at raw PDF bytes decoded as if they were UTF-8 text.
        replacement_ratio = decoded.count("\ufffd") / max(len(decoded), 1)
        if replacement_ratio > 0.3:
            return None
        if len(decoded) > 100:
            return decoded
    except Exception:
        pass
    return None


async def fetch_via_ladder(url: str) -> tuple[list[RetrievalAttempt], str | None]:
    """
    Pure-code fetch ladder. Tries /source → /source.pdf → HTML.
    Returns (attempts, text_or_None). No LLM involved.
    """
    attempts: list[RetrievalAttempt] = []
    norm = normalize_akn_url(url)

    # Ladder 1: /source (PDF/DOCX)
    src = source_url(norm)
    try:
        status, body, ct = await _http_get(src, as_bytes=True)
        attempts.append(RetrievalAttempt(method="/source", url=src,
                        status="success" if status == 200 else f"http_{status}",
                        http_status=status, content_type=ct))
        if status == 200 and isinstance(body, bytes) and len(body) > 500:
            text = _parse_bytes(body, ct)
            if text:
                return attempts, text
    except Exception as e:
        attempts.append(RetrievalAttempt(method="/source", url=src, status="error", error=str(e)[:200]))

    # Ladder 2: /source.pdf
    pdf_url = source_pdf_url(norm)
    try:
        status, body, ct = await _http_get(pdf_url, as_bytes=True)
        attempts.append(RetrievalAttempt(method="/source.pdf", url=pdf_url,
                        status="success" if status == 200 else f"http_{status}",
                        http_status=status, content_type=ct))
        if status == 200 and isinstance(body, bytes) and len(body) > 500:
            text = _parse_bytes(body, ct)
            if text:
                return attempts, text
    except Exception as e:
        attempts.append(RetrievalAttempt(method="/source.pdf", url=pdf_url, status="error", error=str(e)[:200]))

    # Ladder 3: HTML
    if not norm.lower().endswith(".pdf") and not norm.lower().endswith(".docx"):
        try:
            status, html, ct = await _http_get(norm)
            attempts.append(RetrievalAttempt(method="html", url=norm,
                            status="success" if status == 200 else f"http_{status}",
                            http_status=status, content_type=ct))
            if status == 200:
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(" ", strip=True)
                if len(text) > 200:
                    return attempts, text
        except Exception as e:
            attempts.append(RetrievalAttempt(method="html", url=norm, status="error", error=str(e)[:200]))

    return attempts, None


# ===========================================================================
# JSON extractor (pure code)
# ===========================================================================

def _extract_json(text: str) -> dict:
    m = re.search(r'\{.*\}', text or "", re.DOTALL)
    if not m:
        return {}
    try:
        return json.loads(m.group(0))
    except json.JSONDecodeError:
        return {}


# ===========================================================================
# LLM prompts — the LLM makes every decision here
# ===========================================================================

_PROMPT_CLASSIFY = """\
You are a Kenya Law research assistant. Given a user query, decide:

1. document_type: "legislation" | "case_law". If ambiguous, prefer case_law.
2. brave_search_query: the EXACT search query to pass to Brave Search, including
   site: restrictions. Examples:
   - For Employment Act: '"Employment Act" site:new.kenyalaw.org/akn/ke/act'
   - For Rono v Rono: '"Rono" "Rono" site:new.kenyalaw.org/akn/ke/judgment'
   - For a citation: '"[2024] KECA 523" site:new.kenyalaw.org'
   - For issue search: '(procedural fairness termination) site:new.kenyalaw.org/akn/ke/judgment'
3. short_title: clean statute name or null
4. parties: list of party names or empty list
5. neutral_citation: if present or null
6. section: section/article number or null
7. section_label: "Section" or "Article" or null
8. issue_terms: key legal topics
9. is_vague: true if query is too vague to search directly
10. context_query: if is_vague, a general Brave query to clarify (no site restriction)

Return JSON only.
"""

_PROMPT_EVALUATE = """\
You are a Kenya Law candidate evaluator. Given a search result, decide whether
this is the document the user needs.

User query: {user_query}
Expected kind: {expected_kind}

Candidate title: {title}
Candidate URL: {url}
Candidate snippet: {snippet}

Return JSON:
{{"score": 0.0-1.0, "reason": "why this score", "is_relevant": true/false,
 "is_akn_document": true/false}}
"""

_PROMPT_EXTRACT = """\
You are a Kenya Law document parser. Read the beginning of this legal document
and extract metadata.

Document start (first 4000 chars):
{document_start}

Return JSON:
{{"title": "...", "neutral_citation": null, "case_number": null, "court": null,
 "date": null, "parties": [], "act_name": null, "cap": null, "judges": null}}
"""

_PROMPT_VERIFY = """\
You are a Kenya Law document verifier. Given a user query and a candidate document,
decide if this is the right document.

User query: {user_query}
Expected kind: {expected_kind}
Document title: {title}
Document URL: {url}
Document starting text: {doc_start}

Return JSON:
{{"matches": true/false, "score": 0.0-1.0, "reason": "..."}}
"""

_PROMPT_EXTRACT_PROVISION = """\
You are a Kenya Law provision extractor. Find section/article {number} in this
legal document and quote it verbatim. Include the heading and the full text of
the section.

Return JSON:
{{"found": true/false, "heading": "...", "exact_text": "verbatim quote of the entire section", "confidence": 0.0-1.0}}
"""

_PROMPT_EXTRACT_PASSAGES = """\
You are a Kenya Law passage extractor. Find the most relevant passages in this
document related to: {issue_terms}.

Return up to 3 passages. For each, give the exact text and what issue term it matches.

Return JSON:
{{"passages": [{{"text": "...", "matched_term": "...", "relevance": 0.0-1.0}}]}}
"""

_PROMPT_CONTEXT_CLARIFY = """\
You are a Kenya Law research assistant. A user query was too vague to search
directly. Here is context from a general web search:

Web context: {context}

Original query: {original_query}

Now re-classify based on this context. Return the same format as the classification step.
"""


# ===========================================================================
# The tool — code enables, LLM decides
# ===========================================================================

def _fmt(template: str, **kwargs: Any) -> str:
    """Safe format — escapes curly braces in values so .format() never crashes."""
    escaped = {k: str(v).replace("{", "{{").replace("}", "}}") for k, v in kwargs.items()}
    return template.format(**escaped)

class KenyaLawResearchTool:
    """
    Kenya Law research pipeline. Code enables. LLM decides.

    The LLM decides: what to search for, which results are relevant,
    what metadata to extract, what sections/passages to return.

    Code does: HTTP requests, PDF/DOCX byte parsing, JSON cleanup,
    URL normalization, observable recording.
    """

    def __init__(self, llm=None, brave=None) -> None:
        self._llm = llm       # async callable (messages) → str
        self._brave = brave    # BraveSearchClient
        self._audit: list[dict] = []

    async def _llm_call(self, system: str, user: str) -> dict:
        """Call the LLM, parse JSON, log to audit trail."""
        messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
        raw = ""
        try:
            raw = await self._llm.complete(messages, response_format="json", max_tokens=500) if self._llm else "{}"
        except Exception as e:
            raw = '{"error": "' + str(e).replace('"', "'") + '"}'
        parsed = _extract_json(raw)
        self._audit.append({
            "system": system[:200],
            "user": user[:300],
            "raw": raw[:500],
            "parsed": parsed,
        })
        return parsed

    async def research(self, request: KenyaLawRequest) -> KenyaLawResult:
        self._audit = []
        user_query = request.query

        # ─── STEP 1: LLM classifies the query ───────────────────────────────
        # The LLM decides: what type of document, the exact Brave search string
        classification = await self._llm_call(_PROMPT_CLASSIFY, user_query)

        doc_type = classification.get("document_type") or "case_law"
        if request.kind and request.kind != "auto":
            doc_type = request.kind  # agent can override

        search_query = classification.get("brave_search_query", user_query)
        short_title = classification.get("short_title")
        parties = classification.get("parties", [])
        section = str(request.section or classification.get("section") or "")
        issue_terms = classification.get("issue_terms", [])
        is_vague = classification.get("is_vague", False)

        # ─── STEP 2: If vague → clarify via general Brave ───────────────────
        if is_vague and self._brave:
            ctx_query = classification.get("context_query") or user_query
            try:
                ctx_results = await self._brave.search(ctx_query)
                ctx_snippets = []
                for r in ctx_results[:5]:
                    t = getattr(r, "title", "") or ""
                    s = getattr(r, "snippet", "") or ""
                    if t.strip() or s.strip():
                        ctx_snippets.append(f"{t}: {s[:300]}")
                if ctx_snippets:
                    reclass = await self._llm_call(
                        _fmt(_PROMPT_CONTEXT_CLARIFY,
                             context="\n".join(ctx_snippets[:3]),
                             original_query=user_query),
                        "Re-classify with context",
                    )
                    doc_type = reclass.get("document_type", doc_type)
                    search_query = reclass.get("brave_search_query", search_query)
                    short_title = reclass.get("short_title") or short_title
                    parties = reclass.get("parties", []) or parties
                    section = str(reclass.get("section") or "") or section
                    issue_terms = reclass.get("issue_terms", []) or issue_terms
            except Exception:
                pass  # context gathering is best-effort

        # ─── STEP 3: Execute the search the LLM chose ───────────────────────
        # Code enables. The LLM decided the exact search string.
        candidates: list[KenyaLawCandidate] = []
        if self._brave:
            try:
                brave_results = await self._brave.search(search_query)
                for i, r in enumerate(brave_results[:request.max_candidates]):
                    title = (getattr(r, "title", "") or "")[:200]
                    url = (getattr(r, "url", "") or "")[:500]
                    snippet = (getattr(r, "snippet", "") or "")[:300]
                    if url:
                        candidates.append(KenyaLawCandidate(
                            title=title,
                            url=url,
                            snippet=snippet,
                            kind=doc_type,
                            discovery_method="brave",
                        ))
            except Exception:
                pass

        # ─── STEP 4: LLM evaluates each candidate ──────────────────────────
        # The LLM decides which candidates are worth reading.
        ranked: list[tuple[float, KenyaLawCandidate]] = []
        for c in candidates:
            eval_result = await self._llm_call(
                _fmt(_PROMPT_EVALUATE,
                     user_query=user_query,
                     expected_kind=doc_type,
                     title=c.title,
                     url=c.url,
                     snippet=c.snippet or ""),
                f"Evaluate: {c.title}",
            )
            score = eval_result.get("score", 0.3)
            c.score = score
            if eval_result.get("reason"):
                c.reasons = [eval_result.get("reason", "")]
            ranked.append((score, c))

        ranked.sort(key=lambda x: x[0], reverse=True)
        ranked_candidates = [c for _, c in ranked]

        # ─── STEP 5: Fetch the best candidate(s) via read ladder ────────────
        # Pure code. No LLM. Code just gets bytes and turns them into text.
        rejected = []
        all_attempts = []
        selected_doc = None

        for candidate in ranked_candidates[:request.read_top_n]:
            if not candidate.url:
                continue

            fetch_url = candidate.url
            # LLM already told us which URL. Code just fetches it.
            attempts, text = await fetch_via_ladder(fetch_url)
            all_attempts.extend(attempts)

            if text is None or len(text.strip()) < 100:
                rejected.append({"candidate": candidate.title, "reason": "Could not read — all fetch methods failed"})
                continue

            # ─── STEP 6: LLM extracts metadata ─────────────────────────────
            meta = await self._llm_call(
                _fmt(_PROMPT_EXTRACT, document_start=text[:4000]),
                f"Extract metadata from {candidate.title}",
            )
            doc_title = meta.get("title") or candidate.title

            # ─── STEP 7: LLM verifies the document ──────────────────────────
            verify = await self._llm_call(
                _fmt(_PROMPT_VERIFY,
                     user_query=user_query,
                     expected_kind=doc_type,
                     title=doc_title,
                     url=fetch_url,
                     doc_start=text[:2000]),
                f"Verify: {doc_title}",
            )
            if not verify.get("matches", False) and verify.get("score", 0) < 0.3:
                rejected.append({"candidate": candidate.title, "reason": verify.get("reason", "Low relevance")})
                continue

            # ─── DOCUMENT ACCEPTED ─────────────────────────────────────────
            # LLM extracts provision if section requested
            provision_text = None
            if section:
                prov = await self._llm_call(
                    _fmt(_PROMPT_EXTRACT_PROVISION, number=section),
                    text,
                )
                if prov.get("found"):
                    provision_text = prov.get("exact_text", "")

            # LLM extracts passages if issue terms exist
            passages = []
            if issue_terms:
                pas = await self._llm_call(
                    _fmt(_PROMPT_EXTRACT_PASSAGES, issue_terms=", ".join(issue_terms[:5])),
                    text,
                )
                for p in pas.get("passages", []):
                    passages.append(Passage(
                        text=p.get("text", ""),
                        matched_terms=[p.get("matched_term", "")],
                        score=p.get("relevance", 0.5),
                        start_char=0,
                        end_char=0,
                    ))

            selected_doc = KenyaLawDocument(
                kind=doc_type,
                title=doc_title,
                original_url=fetch_url,
                normalized_akn_url=normalize_akn_url(fetch_url),
                content_url=source_url(normalize_akn_url(fetch_url)),
                metadata=meta,
                full_text=text,
                requested_provision_text=provision_text,
                relevant_passages=passages,
                extraction_quality="good" if len(text) > 500 else "minimal",
            )
            break

        # Build result with full text
        selected_text = selected_doc.full_text if selected_doc else None

        # ─── STEP 8: Build result ──────────────────────────────────────────
        suggestions: list[str] = []
        if selected_doc:
            status = "success"
            msg = f"Found and read: {selected_doc.title}"
            if selected_doc.requested_provision_text:
                status = "success_with_provision"
                msg += " — extracted requested provision"
            if selected_doc.relevant_passages:
                status = "success_with_passages"
                msg += f" — extracted {len(selected_doc.relevant_passages)} relevant passage(s)"
        elif rejected:
            status = "candidates_read_no_match"
            msg = f"Read {len(rejected)} candidate(s), none matched."
            suggestions = ["Try a different search strategy.", "Use brave_search for broader discovery."]
        elif candidates:
            status = "candidates_found_not_read"
            msg = f"Found {len(candidates)} candidate(s) but none could be read."
            suggestions = ["Try direct URL fetch with general_web_fetch."]
        else:
            status = "no_candidates"
            msg = "No candidates found on Kenya Law."
            suggestions = ["Try brave_search for broader discovery.", "Clarify the query."]

        return KenyaLawResult(
            status=status,
            message=msg,
            query=request.query,
            candidates=ranked_candidates,
            rejected_candidates=rejected,
            retrieval_attempts=all_attempts,
            selected_document_title=selected_doc.title if selected_doc else None,
            selected_document_url=selected_doc.normalized_akn_url if selected_doc else None,
            selected_document_provision=selected_doc.requested_provision_text if selected_doc else None,
            selected_document_text=selected_text,
            suggested_next_actions=suggestions,
            llm_audit=list(self._audit),
        )
