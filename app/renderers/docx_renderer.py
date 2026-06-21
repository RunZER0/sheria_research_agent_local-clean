import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.document_spec import DocumentSpec


class DocxRenderer:
    @staticmethod
    def render(spec: DocumentSpec, output_path: str) -> str:
        doc = Document()

        # 1. Apply Layout Style Specs
        sections = doc.sections
        for section in sections:
            if spec.style.margins == "narrow":
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            else:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

        # 2. Iterate Blocks and Hydrate Document Elements
        for block in spec.blocks:
            if block.type == "heading":
                p = doc.add_heading(block.text, level=block.level)
                p.style.font.name = spec.style.font_name
                p.style.font.size = Pt(spec.style.font_size + 2)

            elif block.type == "paragraph":
                p = doc.add_paragraph(block.text)
                p.paragraph_format.line_spacing = spec.style.line_spacing
                if spec.style.alignment == "justify":
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = spec.style.font_name
                    run.font.size = Pt(spec.style.font_size)

            elif block.type == "table":
                table_el = doc.add_table(rows=1, cols=len(block.columns))
                hdr_cells = table_el.rows[0].cells
                for i, col_name in enumerate(block.columns):
                    hdr_cells[i].text = col_name
                for row_data in block.rows:
                    row_cells = table_el.add_row().cells
                    for i, cell_value in enumerate(row_data):
                        row_cells[i].text = str(cell_value)

            elif block.type == "page_break":
                doc.add_page_break()

        doc.save(output_path)
        return output_path
