from __future__ import annotations

import mimetypes
import re
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Callable

import httpx

from app.schemas.research_state import FetchResult


_HTTP_USER_AGENT = (
    "Mozilla/5.0 (X11; Linux x86_64; rv:121.0) "
    "Gecko/20100101 Firefox/121.0"
)


def looks_like_xml(text: str) -> bool:
    sample = text.lstrip()[:500].lower()
    return sample.startswith("<?xml") or "<akoma" in sample or "<akn" in sample


def looks_like_html(text: str) -> bool:
    sample = text.lstrip()[:500].lower()
    return "<html" in sample or "<!doctype html" in sample


def strip_html(html: str) -> str:
    html = re.sub(r"(?is)<script.*?>.*?</script>", " ", html)
    html = re.sub(r"(?is)<style.*?>.*?</style>", " ", html)
    text = re.sub(r"(?s)<[^>]+>", " ", html)
    return re.sub(r"\s+", " ", text).strip()


def strip_xml(xml: str) -> str:
    text = re.sub(r"(?s)<[^>]+>", " ", xml)
    return re.sub(r"\s+", " ", text).strip()


class FetchManager:
    """Manages fetching and ingesting source documents with a recovery ladder.

    The read ladder when primary fetch fails:
      1. httpx fetch with browser-like User-Agent
      2. If 403, try browser fetch via Playwright Firefox
      3. If browser fails, try alternative URLs (PDF download, old Kenya Law mirrors)
      4. If all attempts fail, return FetchResult(ok=False)
    """

    def __init__(self, timeout: int = 20) -> None:
        self.timeout = timeout

    async def fetch_and_ingest(
        self,
        url: str,
        expected_title: str | None = None,
        alternative_urls: list[str] | None = None,
    ) -> FetchResult:
        """Primary fetch — httpx with fallback to browser and alternative URLs."""
        # Step 1: try httpx
        result = await self._fetch_httpx(url, expected_title)
        if result.ok:
            return result

        # Step 2: try browser fetch for 403 (site protection)
        if result.status_code == 403:
            browser_result = await self._fetch_browser(url, expected_title)
            if browser_result.ok:
                return browser_result

        # Step 3: try alternative URLs (PDF downloads, mirrors)
        if alternative_urls:
            for alt_url in alternative_urls:
                alt_result = await self._fetch_httpx(alt_url, expected_title)
                if alt_result.ok:
                    return alt_result
                if alt_result.status_code == 403:
                    browser_alt = await self._fetch_browser(alt_url, expected_title)
                    if browser_alt.ok:
                        return browser_alt

        # Step 4: return the original failure result
        return result

    async def _fetch_httpx(self, url: str, expected_title: str | None = None) -> FetchResult:
        """Standard httpx fetch."""
        try:
            async with httpx.AsyncClient(
                timeout=httpx.Timeout(self.timeout),
                follow_redirects=True,
            ) as client:
                response = await client.get(
                    url,
                    headers={"User-Agent": _HTTP_USER_AGENT},
                )
                content_type = response.headers.get("content-type", "")
                final_url = str(response.url)
                status = response.status_code
                raw = response.content

                if status != 200:
                    return FetchResult(
                        ok=False,
                        url=url,
                        final_url=final_url,
                        status_code=status,
                        fetch_method="http_fetch",
                        parser_used="none",
                        error=f"HTTP {status}",
                    )

            parser_used = self._choose_parser(url, content_type, raw)
            text = self._parse(raw, parser_used)

            return FetchResult(
                ok=bool(text.strip()),
                url=url,
                final_url=final_url,
                status_code=status,
                content_type=content_type,
                parser_used=parser_used,
                fetch_method="http_fetch",
                title=expected_title or final_url or url,
                text=text,
                metadata={"bytes": len(raw)},
                error=None if text.strip() else "Fetched source but extracted no text.",
            )

        except Exception as exc:
            return FetchResult(
                ok=False,
                url=url,
                fetch_method="http_fetch",
                parser_used="none",
                error=str(exc),
            )

    async def _fetch_browser(self, url: str, expected_title: str | None = None) -> FetchResult:
        """Headless Playwright Firefox browser rendering for publicly accessible pages."""
        try:
            from playwright.async_api import async_playwright

            async with async_playwright() as p:
                browser = await p.firefox.launch(headless=True)
                page = await browser.new_page(user_agent=_HTTP_USER_AGENT)
                await page.goto(url, wait_until="domcontentloaded", timeout=self.timeout * 1000)
                await page.wait_for_timeout(1000)
                html = await page.content()
                final_url = page.url
                await browser.close()

            raw = html.encode("utf-8")
            parser_used = self._choose_parser(url, "text/html", raw)
            text = self._parse(raw, parser_used)

            return FetchResult(
                ok=bool(text.strip()),
                url=url,
                final_url=final_url,
                status_code=200,
                content_type="text/html",
                parser_used=parser_used or "browser_fetch_firefox",
                fetch_method="browser_fetch_firefox",
                title=expected_title or final_url or url,
                text=text,
                metadata={"browser": "firefox"},
                error=None if text.strip() else "Browser rendered page but extracted no text.",
            )

        except ImportError:
            return FetchResult(
                ok=False, url=url, fetch_method="browser_fetch_firefox",
                parser_used="none",
                error="Playwright is not installed. Cannot use browser fetch.",
            )
        except Exception as exc:
            return FetchResult(
                ok=False, url=url, fetch_method="browser_fetch_firefox",
                parser_used="none", error=str(exc),
            )

    def _choose_parser(self, url: str, content_type: str, raw: bytes) -> str:
        lower_url = url.lower()
        lower_ct = content_type.lower()

        if "pdf" in lower_ct or lower_url.endswith(".pdf"):
            return "pdf_text_extract"

        if (
            "wordprocessingml" in lower_ct
            or "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in lower_ct
            or lower_url.endswith(".docx")
        ):
            return "docx_text_extract"

        sample = raw[:2000].decode("utf-8", errors="ignore")

        if "xml" in lower_ct or looks_like_xml(sample):
            return "akn_xml_parser"

        if "html" in lower_ct or looks_like_html(sample):
            return "html_legal_document_parser"

        return "plain_text_parser"

    def _parse(self, raw: bytes, parser_used: str) -> str:
        if parser_used == "html_legal_document_parser":
            return strip_html(raw.decode("utf-8", errors="ignore"))

        if parser_used == "akn_xml_parser":
            return strip_xml(raw.decode("utf-8", errors="ignore"))

        if parser_used == "plain_text_parser":
            return raw.decode("utf-8", errors="ignore").strip()

        if parser_used == "pdf_text_extract":
            return self._parse_pdf(raw)

        if parser_used == "docx_text_extract":
            return self._parse_docx(raw)

        return ""

    def _parse_pdf(self, raw: bytes) -> str:
        try:
            from pypdf import PdfReader
            with NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
                tmp.write(raw)
                tmp.flush()
                reader = PdfReader(tmp.name)
                parts = []
                for page in reader.pages:
                    parts.append(page.extract_text() or "")
                return "\n".join(parts).strip()
        except Exception:
            return ""

    def _parse_docx(self, raw: bytes) -> str:
        try:
            from docx import Document
            with NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
                tmp.write(raw)
                tmp.flush()
                doc = Document(tmp.name)
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        parts.append(" | ".join(cell.text.strip() for cell in row.cells))
                return "\n".join(parts).strip()
        except Exception:
            return ""
