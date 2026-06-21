from __future__ import annotations

import asyncio
import json
import os
import re
import shutil
import subprocess
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Awaitable, Callable, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

from app.schemas.document_export import (
    BulletListBlock,
    ConstraintFailure,
    ConstraintReport,
    DocumentBlock,
    DocumentConstraints,
    DocumentExportRequest,
    DocumentExportResult,
    DocumentSpec,
    DocumentStyle,
    HeadingBlock,
    NumberedListBlock,
    OutputFormat,
    PageBreakBlock,
    PageCountMode,
    ParagraphBlock,
    SourceBasisBlock,
    TableBlock,
    WordCountScope,
)


ARTIFACT_ROOT = Path("data/artifacts/generated")


RevisionCallback = Callable[[DocumentSpec, ConstraintReport], Awaitable[DocumentSpec]]


def _strip_inline_markdown(text: str) -> str:
    """Remove inline markdown formatting markers from text."""
    # Remove **bold** and __bold__
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # Remove *italic* and _italic_
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
    text = re.sub(r'(?<!_)_(?!_)(.+?)(?<!_)_(?!_)', r'\1', text)
    # Remove `code`
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove ~~strikethrough~~
    text = re.sub(r'~~(.+?)~~', r'\1', text)
    return text


def count_words(text: str) -> int:
    return len(re.findall(r"\b[\w]+(?:[''\-][\w]+)?\b", text, flags=re.UNICODE))


def split_answer_to_blocks(answer_text: str) -> list[DocumentBlock]:
    """
    Converts a normal answer into a usable block structure.

    Supported lightweight syntax:
    - Markdown headings: #, ##, ###
    - Markdown tables
    - Bullet lists beginning with "- "
    - Numbered lists beginning with "1. "
    """
    lines = answer_text.strip().splitlines()
    blocks: list[DocumentBlock] = []
    paragraph_buffer: list[str] = []
    i = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        text = "\n".join(paragraph_buffer).strip()
        if text:
            blocks.append(
                ParagraphBlock(
                    id=f"p_{len(blocks)+1}",
                    text=_strip_inline_markdown(text),
                )
            )
        paragraph_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        if line.startswith("#"):
            flush_paragraph()
            level = min(len(line) - len(line.lstrip("#")), 4)
            text = _strip_inline_markdown(line.lstrip("#").strip())
            blocks.append(
                HeadingBlock(
                    id=f"h_{len(blocks)+1}",
                    level=level,
                    text=text,
                )
            )
            i += 1
            continue

        # Markdown table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}:?\s*\|", lines[i + 1]):
            flush_paragraph()
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows: list[list[str]] = []

            while i < len(lines) and "|" in lines[i]:
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1

            blocks.append(
                TableBlock(
                    id=f"t_{len(blocks)+1}",
                    columns=header,
                    rows=rows,
                )
            )
            continue

        # Bullet list
        if line.strip().startswith("- "):
            flush_paragraph()
            items: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(_strip_inline_markdown(lines[i].strip()[2:].strip()))
                i += 1
            blocks.append(
                BulletListBlock(
                    id=f"b_{len(blocks)+1}",
                    items=items,
                )
            )
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            flush_paragraph()
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(_strip_inline_markdown(re.sub(r"^\s*\d+\.\s+", "", lines[i]).strip()))
                i += 1
            blocks.append(
                NumberedListBlock(
                    id=f"n_{len(blocks)+1}",
                    items=items,
                )
            )
            continue

        paragraph_buffer.append(line)
        i += 1

    flush_paragraph()

    if not blocks:
        blocks.append(ParagraphBlock(id="p_1", text=answer_text.strip()))

    return blocks


def extract_countable_text(spec: DocumentSpec, scope: WordCountScope) -> str:
    pieces: list[str] = []

    for block in spec.blocks:
        if isinstance(block, HeadingBlock):
            if scope == WordCountScope.entire_document:
                pieces.append(block.text)

        elif isinstance(block, ParagraphBlock):
            pieces.append(block.text)

        elif isinstance(block, BulletListBlock):
            pieces.extend(block.items)

        elif isinstance(block, NumberedListBlock):
            pieces.extend(block.items)

        elif isinstance(block, TableBlock):
            if scope in (WordCountScope.include_tables, WordCountScope.entire_document):
                if block.caption:
                    pieces.append(block.caption)
                pieces.extend(block.columns)
                for row in block.rows:
                    pieces.extend(row)

        elif isinstance(block, SourceBasisBlock):
            if scope == WordCountScope.entire_document:
                for item in block.items:
                    pieces.append(f"{item.source} {item.role} {item.strength}")

    return "\n".join(pieces)


def estimate_words_for_pages(style: DocumentStyle, pages: int) -> int:
    """
    Rough starting estimate only. Actual page count is verified after rendering.
    Times New Roman 12, 1.5 spacing, normal margins is roughly 330-420 words/page.
    """
    base = 375

    if style.line_spacing >= 1.5:
        base = 360
    if style.line_spacing >= 2:
        base = 275
    if style.font_size_pt > 12:
        base -= (style.font_size_pt - 12) * 25
    if style.font_size_pt < 12:
        base += (12 - style.font_size_pt) * 25

    return max(150, base * pages)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


class WordCountVerifier:
    def verify(self, spec: DocumentSpec) -> ConstraintReport:
        constraint = spec.constraints.word_count
        if not constraint:
            return ConstraintReport(ok=True)

        text = extract_countable_text(spec, constraint.scope)
        actual = count_words(text)
        target = constraint.target
        tolerance = constraint.tolerance

        if abs(actual - target) <= tolerance:
            return ConstraintReport(ok=True)

        if actual < target:
            return ConstraintReport(
                ok=False,
                failures=[
                    ConstraintFailure(
                        code="WORD_COUNT_TOO_LOW",
                        message=f"Word count is {actual}; target is {target}.",
                        expected=target,
                        actual=actual,
                        suggested_action=f"Add exactly {target - actual} words to body text.",
                    )
                ],
            )

        return ConstraintReport(
            ok=False,
            failures=[
                ConstraintFailure(
                    code="WORD_COUNT_TOO_HIGH",
                    message=f"Word count is {actual}; target is {target}.",
                    expected=target,
                    actual=actual,
                    suggested_action=f"Remove exactly {actual - target} words from body text.",
                )
            ],
        )


class DeterministicRevisionService:
    """
    Fallback revision service.

    This does not replace your LLM RevisionAgent. It gives the system a working
    deterministic correction path for small exact word-count gaps.
    """

    filler_words = [
        "Accordingly",
        "the",
        "analysis",
        "remains",
        "grounded",
        "in",
        "the",
        "available",
        "legal",
        "materials",
        "and",
        "should",
        "be",
        "read",
        "with",
        "the",
        "stated",
        "source",
        "limitations",
        "carefully",
    ]

    def revise(self, spec: DocumentSpec, report: ConstraintReport) -> DocumentSpec:
        if not report.failures:
            return spec

        failure = report.failures[0]

        if failure.code == "WORD_COUNT_TOO_LOW":
            missing = int(failure.expected) - int(failure.actual)
            self._add_words(spec, missing)

        elif failure.code == "WORD_COUNT_TOO_HIGH":
            excess = int(failure.actual) - int(failure.expected)
            self._remove_words(spec, excess)

        return spec

    def _last_paragraph(self, spec: DocumentSpec) -> Optional[ParagraphBlock]:
        for block in reversed(spec.blocks):
            if isinstance(block, ParagraphBlock):
                return block
        return None

    def _add_words(self, spec: DocumentSpec, n: int) -> None:
        if n <= 0:
            return

        paragraph = self._last_paragraph(spec)
        if not paragraph:
            paragraph = ParagraphBlock(id=f"p_{len(spec.blocks)+1}", text="")
            spec.blocks.append(paragraph)

        words = []
        while len(words) < n:
            words.extend(self.filler_words)

        addition = " ".join(words[:n])
        paragraph.text = paragraph.text.rstrip() + " " + addition + "."

    def _remove_words(self, spec: DocumentSpec, n: int) -> None:
        if n <= 0:
            return

        paragraph = self._last_paragraph(spec)
        if not paragraph:
            return

        words = re.findall(r"\b[\w]+(?:[''\-][\w]+)?\b|[^\w\s]", paragraph.text, flags=re.UNICODE)
        word_indexes = [
            idx for idx, token in enumerate(words)
            if re.match(r"\b[\w]+(?:[''\-][\w]+)?\b", token, flags=re.UNICODE)
        ]

        remove_indexes = set(word_indexes[-n:])
        kept = [token for idx, token in enumerate(words) if idx not in remove_indexes]

        text = " ".join(kept)
        text = re.sub(r"\s+([,.;:!?])", r"\1", text)
        paragraph.text = text.strip()


class DocxRenderer:
    def render(self, spec: DocumentSpec, out_path: Path) -> Path:
        out_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._apply_style(doc, spec.style)

        doc.core_properties.title = spec.title

        title = doc.add_paragraph()
        title_run = title.add_run(spec.title)
        title_run.bold = True
        title_run.font.name = spec.style.font_family
        title_run.font.size = Pt(spec.style.font_size_pt + 2)

        for block in spec.blocks:
            self._render_block(doc, block, spec.style)

        doc.save(out_path)
        return out_path

    def _apply_style(self, doc: Document, style: DocumentStyle) -> None:
        section = doc.sections[0]

        if style.page_size == "A4":
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        else:
            section.page_width = Cm(21.59)
            section.page_height = Cm(27.94)

        section.top_margin = Cm(style.margin_top_cm)
        section.bottom_margin = Cm(style.margin_bottom_cm)
        section.left_margin = Cm(style.margin_left_cm)
        section.right_margin = Cm(style.margin_right_cm)

        normal = doc.styles["Normal"]
        normal.font.name = style.font_family
        normal.font.size = Pt(style.font_size_pt)

        paragraph_format = normal.paragraph_format
        paragraph_format.line_spacing = style.line_spacing
        paragraph_format.space_after = Pt(6)

    def _apply_run_style(self, run, style: DocumentStyle, bold: bool = False) -> None:
        run.font.name = style.font_family
        run._element.rPr.rFonts.set(qn("w:eastAsia"), style.font_family)
        run.font.size = Pt(style.font_size_pt)
        run.bold = bold

    def _render_block(self, doc: Document, block: DocumentBlock, style: DocumentStyle) -> None:
        if isinstance(block, HeadingBlock):
            p = doc.add_heading(level=block.level)
            run = p.add_run(block.text)
            self._apply_run_style(run, style, bold=True)

        elif isinstance(block, ParagraphBlock):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = style.line_spacing
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(block.text)
            self._apply_run_style(run, style)

            if block.citations:
                cite_run = p.add_run(" " + " ".join(block.citations))
                self._apply_run_style(cite_run, style)
                cite_run.italic = True

        elif isinstance(block, BulletListBlock):
            for item in block.items:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing = style.line_spacing
                run = p.add_run(item)
                self._apply_run_style(run, style)

        elif isinstance(block, NumberedListBlock):
            for item in block.items:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.line_spacing = style.line_spacing
                run = p.add_run(item)
                self._apply_run_style(run, style)

        elif isinstance(block, TableBlock):
            if block.caption:
                p = doc.add_paragraph()
                run = p.add_run(block.caption)
                self._apply_run_style(run, style, bold=True)

            table = doc.add_table(rows=1, cols=len(block.columns))
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            for idx, column in enumerate(block.columns):
                hdr_cells[idx].text = column
                set_cell_shading(hdr_cells[idx], "D9EAF7")

            for row in block.rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row[: len(block.columns)]):
                    cells[idx].text = value

            doc.add_paragraph()

        elif isinstance(block, PageBreakBlock):
            doc.add_page_break()

        elif isinstance(block, SourceBasisBlock):
            doc.add_heading("Source Basis", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Source"
            table.rows[0].cells[1].text = "Role"
            table.rows[0].cells[2].text = "Strength"

            for item in block.items:
                row = table.add_row().cells
                row[0].text = item.source
                row[1].text = item.role
                row[2].text = item.strength


class PdfPublisher:
    def convert_docx_to_pdf(self, docx_path: Path, output_dir: Path) -> Optional[Path]:
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            return None

        output_dir.mkdir(parents=True, exist_ok=True)

        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )

        pdf_path = output_dir / f"{docx_path.stem}.pdf"
        return pdf_path if pdf_path.exists() else None


class PdfVerifier:
    def page_count(self, pdf_path: Path) -> Optional[int]:
        if not PdfReader:
            return None

        try:
            reader = PdfReader(str(pdf_path))
            return len(reader.pages)
        except Exception:
            return None

    def page_text(self, pdf_path: Path, page_number: int) -> str:
        if not PdfReader:
            return ""

        try:
            reader = PdfReader(str(pdf_path))
            page = reader.pages[page_number - 1]
            return page.extract_text() or ""
        except Exception:
            return ""

    def verify(self, pdf_path: Path, spec: DocumentSpec) -> ConstraintReport:
        failures: list[ConstraintFailure] = []

        page_count_constraint = spec.constraints.page_count
        actual_pages = self.page_count(pdf_path)

        if page_count_constraint and actual_pages is not None:
            expected = page_count_constraint.value

            if page_count_constraint.mode == PageCountMode.exact and actual_pages != expected:
                failures.append(
                    ConstraintFailure(
                        code="PAGE_COUNT_MISMATCH",
                        message=f"PDF has {actual_pages} pages; expected exactly {expected}.",
                        expected=expected,
                        actual=actual_pages,
                        suggested_action="Revise content length or insert/remove page breaks.",
                    )
                )

            elif page_count_constraint.mode == PageCountMode.max and actual_pages > expected:
                failures.append(
                    ConstraintFailure(
                        code="PAGE_COUNT_TOO_HIGH",
                        message=f"PDF has {actual_pages} pages; maximum is {expected}.",
                        expected=expected,
                        actual=actual_pages,
                        suggested_action="Condense body text or tables.",
                    )
                )

            elif page_count_constraint.mode == PageCountMode.min and actual_pages < expected:
                failures.append(
                    ConstraintFailure(
                        code="PAGE_COUNT_TOO_LOW",
                        message=f"PDF has {actual_pages} pages; minimum is {expected}.",
                        expected=expected,
                        actual=actual_pages,
                        suggested_action="Expand body text.",
                    )
                )

        for req in spec.constraints.page_requirements:
            text = self.page_text(pdf_path, req.page_number)

            for required in req.must_include_text:
                if required.lower() not in text.lower():
                    failures.append(
                        ConstraintFailure(
                            code="PAGE_TEXT_MISSING",
                            message=f"Page {req.page_number} does not include required text: {required}",
                            expected=required,
                            actual=text[:500],
                            suggested_action=f"Move or add required text to page {req.page_number}.",
                        )
                    )

        return ConstraintReport(ok=not failures, failures=failures)


class PageCountRevisionService:
    """
    Simple deterministic page-count adjustment fallback.

    For serious legal text quality, replace this with an LLM revision callback.
    """

    def revise(self, spec: DocumentSpec, report: ConstraintReport) -> DocumentSpec:
        if not report.failures:
            return spec

        failure = report.failures[0]

        if failure.code == "PAGE_COUNT_TOO_LOW":
            self._expand(spec)
        elif failure.code in {"PAGE_COUNT_TOO_HIGH", "PAGE_COUNT_MISMATCH"}:
            if isinstance(failure.actual, int) and isinstance(failure.expected, int):
                if failure.actual > failure.expected:
                    self._condense(spec)
                else:
                    self._expand(spec)

        return spec

    def _last_paragraph(self, spec: DocumentSpec) -> Optional[ParagraphBlock]:
        for block in reversed(spec.blocks):
            if isinstance(block, ParagraphBlock):
                return block
        return None

    def _expand(self, spec: DocumentSpec) -> None:
        paragraph = self._last_paragraph(spec)
        if paragraph:
            paragraph.text += (
                "\n\nThis additional analysis explains the practical implications, "
                "the limits of the available sources, and the reasons the conclusion "
                "should be treated as grounded but subject to further verification "
                "where fuller records or newer authorities become available."
            )

    def _condense(self, spec: DocumentSpec) -> None:
        for block in spec.blocks:
            if isinstance(block, ParagraphBlock):
                sentences = re.split(r"(?<=[.!?])\s+", block.text.strip())
                if len(sentences) > 2:
                    block.text = " ".join(sentences[:-1])
                    return


class DocumentExportService:
    def __init__(
        self,
        artifact_root: Path = ARTIFACT_ROOT,
        revision_callback: Optional[RevisionCallback] = None,
    ) -> None:
        self.artifact_root = artifact_root
        self.word_verifier = WordCountVerifier()
        self.docx_renderer = DocxRenderer()
        self.pdf_publisher = PdfPublisher()
        self.pdf_verifier = PdfVerifier()
        self.word_reviser = DeterministicRevisionService()
        self.page_reviser = PageCountRevisionService()
        self.revision_callback = revision_callback

    def build_spec_from_request(self, request: DocumentExportRequest) -> DocumentSpec:
        blocks = split_answer_to_blocks(request.answer_text)

        return DocumentSpec(
            title=request.title,
            document_type="sheria_generated_document",
            output_formats=request.output_formats,
            style=request.style,
            constraints=request.constraints,
            blocks=blocks,
            metadata={
                "user_request": request.user_request,
            },
        )

    async def export(self, request: DocumentExportRequest, max_attempts: int = 6) -> DocumentExportResult:
        # Generate contextual filename from the document title
        safe_name = re.sub(r'[^\w\s-]', '', request.title)
        safe_name = re.sub(r'\s+', '_', safe_name.strip())[:80].lower() or "sheria_document"
        slug = f"{safe_name}_{uuid.uuid4().hex[:6]}"
        out_dir = self.artifact_root / slug
        out_dir.mkdir(parents=True, exist_ok=True)

        spec = self.build_spec_from_request(request)
        warnings: list[str] = []
        final_report = ConstraintReport(ok=True)

        docx_path = out_dir / f"{safe_name}.docx"
        pdf_path: Optional[Path] = None

        for attempt in range(1, max_attempts + 1):
            word_report = self.word_verifier.verify(spec)

            if not word_report.ok:
                final_report = word_report
                spec = await self._revise(spec, word_report)
                continue

            self.docx_renderer.render(spec, docx_path)

            if OutputFormat.pdf in spec.output_formats or spec.constraints.page_count or spec.constraints.page_requirements:
                converted = self.pdf_publisher.convert_docx_to_pdf(docx_path, out_dir)

                if not converted:
                    warnings.append(
                        "PDF conversion requires LibreOffice/soffice installed on the backend."
                    )
                    final_report = ConstraintReport(
                        ok=False,
                        failures=[
                            ConstraintFailure(
                                code="PDF_CONVERSION_UNAVAILABLE",
                                message="LibreOffice/soffice was not found, so PDF/page verification could not run.",
                            )
                        ],
                    )
                    break

                pdf_path = converted
                pdf_report = self.pdf_verifier.verify(pdf_path, spec)

                if not pdf_report.ok:
                    final_report = pdf_report
                    spec = await self._revise(spec, pdf_report)
                    continue

            final_report = ConstraintReport(ok=True)
            break

        manifest_path = self._write_manifest(
            out_dir=out_dir,
            document_id=safe_name,
            request=request,
            spec=spec,
            docx_path=docx_path if docx_path.exists() else None,
            pdf_path=pdf_path if pdf_path and pdf_path.exists() else None,
            report=final_report,
            warnings=warnings,
        )

        return DocumentExportResult(
            ok=final_report.ok,
            document_id=safe_name,
            docx_path=str(docx_path) if docx_path.exists() else None,
            pdf_path=str(pdf_path) if pdf_path and pdf_path.exists() else None,
            artifact_manifest_path=str(manifest_path),
            constraint_report=final_report,
            warnings=warnings,
        )

    async def _revise(self, spec: DocumentSpec, report: ConstraintReport) -> DocumentSpec:
        if self.revision_callback:
            return await self.revision_callback(spec, report)

        raise RuntimeError(
            "revision_callback is required when constraints fail. "
            "Provide an LLM-based revision callback to repair content."
        )

    def _write_manifest(
        self,
        out_dir: Path,
        document_id: str,
        request: DocumentExportRequest,
        spec: DocumentSpec,
        docx_path: Optional[Path],
        pdf_path: Optional[Path],
        report: ConstraintReport,
        warnings: list[str],
    ) -> Path:
        manifest_path = out_dir / "artifact_manifest.json"

        body_text = extract_countable_text(spec, WordCountScope.body_only)

        manifest = {
            "document_id": document_id,
            "ok": report.ok,
            "request": request.model_dump(mode="json"),
            "style": spec.style.model_dump(mode="json"),
            "constraints": spec.constraints.model_dump(mode="json"),
            "word_count_body_only": count_words(body_text),
            "docx_path": str(docx_path) if docx_path else None,
            "pdf_path": str(pdf_path) if pdf_path else None,
            "constraint_report": report.model_dump(mode="json"),
            "warnings": warnings,
        }

        manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
        return manifest_path
